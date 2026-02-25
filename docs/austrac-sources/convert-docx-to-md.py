"""Convert all AUSTRAC .docx files to clean Markdown.

Handles: headings, bullet/numbered lists, tables (with bullets inside cells),
instruction blocks, proper paragraph separation after list items,
cross-references (internal bookmark links → markdown anchors),
and external hyperlinks (→ markdown links with URLs).

Usage:
    python convert-docx-to-md.py              # Convert all sectors
    python convert-docx-to-md.py --file X.docx --out Y.md   # Convert one file
"""
import os
import re
import sys
from docx import Document
from lxml import etree

SECTORS = ['real-estate', 'accountants', 'jewellers']
CATEGORIES = ['program-core', 'personnel-forms', 'customer-forms', 'maintain-program-forms']
BASE = os.path.dirname(os.path.abspath(__file__))


# ─── XML NAMESPACES ──────────────────────────────────────────────────────────

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS = {'w': W_NS, 'r': R_NS}


# ─── HYPERLINK & CROSS-REFERENCE HELPERS ─────────────────────────────────────

def _heading_to_anchor(heading_text):
    """Convert heading text to a GitHub-flavoured markdown anchor.

    '## Risk assessment sources' → 'risk-assessment-sources'
    """
    anchor = heading_text.lower().strip()
    anchor = re.sub(r'[^\w\s-]', '', anchor)   # remove punctuation
    anchor = re.sub(r'\s+', '-', anchor)        # spaces → hyphens
    anchor = re.sub(r'-+', '-', anchor)         # collapse multiple hyphens
    return anchor.strip('-')


def _build_bookmark_map(doc):
    """Build a map from Word bookmark names → heading text.

    Scans all bookmarkStart elements and maps their names to the text of
    the paragraph they appear in. This lets us resolve cross-references
    like _Risk_assessment_sources → 'Risk assessment sources'.
    """
    bookmark_map = {}  # bookmark_name → paragraph_text
    body = doc.element.body
    for bm_start in body.findall(f'.//{{{W_NS}}}bookmarkStart'):
        bm_name = bm_start.get(f'{{{W_NS}}}name', '')
        if not bm_name or bm_name == '_GoBack':
            continue
        # Walk up to the parent paragraph
        parent = bm_start.getparent()
        if parent is not None and parent.tag == f'{{{W_NS}}}p':
            para_text = ''.join(
                r.text or '' for r in parent.findall(f'.//{{{W_NS}}}t')
            ).strip()
            if para_text:
                bookmark_map[bm_name] = para_text
    return bookmark_map


def _build_rel_map(doc):
    """Build a map from relationship IDs → external URLs."""
    rel_map = {}
    for rel_id, rel in doc.part.rels.items():
        if rel.is_external:
            rel_map[rel_id] = rel.target_ref
    return rel_map


def _para_to_md_text(para_element, bookmark_map, rel_map):
    """Convert a paragraph's XML element to markdown text with links.

    Walks through runs, hyperlinks, and structured content controls:
    - Internal cross-refs (w:hyperlink with w:anchor) → [text](#anchor)
    - External links (w:hyperlink with r:id) → [text](url)
    - Checkboxes (w:sdt with w14:checkbox) → [ ] or [x]
    - Plain runs → plain text
    """
    return _para_to_md_text_with_checkboxes(para_element, bookmark_map, rel_map)


W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'


def _extract_checkboxes(element):
    """Find w14:checkbox elements and return their checked state.

    Returns a list of (checked: bool) for each checkbox found.
    """
    checkboxes = element.findall(f'.//{{{W14_NS}}}checkbox')
    results = []
    for cb in checkboxes:
        checked_el = cb.find(f'{{{W14_NS}}}checked')
        is_checked = (checked_el is not None
                      and checked_el.get(f'{{{W14_NS}}}val', '0') == '1')
        results.append(is_checked)
    return results


def _para_to_md_text_with_checkboxes(para_element, bookmark_map, rel_map):
    """Convert paragraph XML to markdown text, handling checkboxes.

    Structured content controls (w:sdt) containing w14:checkbox are
    rendered as [ ] (unchecked) or [x] (checked).
    """
    parts = []
    # Field code state: tracks HYPERLINK field codes (fldChar begin/separate/end)
    field_url = None       # URL from instrText, set on 'begin'
    field_text_parts = []  # Display text, collected between 'separate' and 'end'
    in_field = False       # True between 'begin' and 'end'
    in_field_text = False  # True between 'separate' and 'end'

    for child in para_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sdt':
            # Structured content control — check for checkbox
            cbs = _extract_checkboxes(child)
            if cbs:
                # It's a checkbox — render it
                for checked in cbs:
                    parts.append('[x]' if checked else '[ ]')
            else:
                # Not a checkbox sdt — extract its text
                for t in child.findall(f'.//{{{W_NS}}}t'):
                    if t.text:
                        if in_field_text:
                            field_text_parts.append(t.text)
                        else:
                            parts.append(t.text)

        elif tag == 'hyperlink':
            anchor = child.get(f'{{{W_NS}}}anchor', '')
            r_id = child.get(f'{{{R_NS}}}id', '')
            link_text = ''.join(
                r.text or '' for r in child.findall(f'.//{{{W_NS}}}t')
            ).strip()
            if not link_text:
                continue
            if r_id and r_id in rel_map:
                # External URL (may also have a fragment anchor)
                url = rel_map[r_id]
                if anchor:
                    url = f'{url}#{anchor}'
                parts.append(f'[{link_text}]({url})')
            elif anchor:
                # Pure internal cross-reference (no external URL)
                heading_text = bookmark_map.get(anchor, '')
                if heading_text:
                    md_anchor = _heading_to_anchor(heading_text)
                else:
                    md_anchor = anchor.lstrip('_').lower()
                    md_anchor = re.sub(r'[\s_]+', '-', md_anchor)
                    md_anchor = re.sub(r'-+', '-', md_anchor).strip('-')
                parts.append(f'[{link_text}](#{md_anchor})')
            else:
                parts.append(link_text)

        elif tag == 'r':
            # Check for field characters (fldChar) and field instructions
            fldChar = child.find(f'{{{W_NS}}}fldChar')
            instrText = child.find(f'{{{W_NS}}}instrText')

            if fldChar is not None:
                fld_type = fldChar.get(f'{{{W_NS}}}fldCharType', '')
                if fld_type == 'begin':
                    in_field = True
                    field_url = None
                    field_text_parts = []
                    in_field_text = False
                elif fld_type == 'separate':
                    in_field_text = True
                elif fld_type == 'end':
                    # Emit the field-code hyperlink
                    if field_url and field_text_parts:
                        link_text = ''.join(field_text_parts).strip()
                        if link_text:
                            parts.append(f'[{link_text}]({field_url})')
                    in_field = False
                    in_field_text = False
                    field_url = None
                    field_text_parts = []
            elif instrText is not None and in_field:
                # Parse HYPERLINK instruction: HYPERLINK "url"
                instr = (instrText.text or '').strip()
                m = re.match(r'HYPERLINK\s+"([^"]+)"', instr)
                if m:
                    field_url = m.group(1)
            else:
                # Regular text run
                for t in child.findall(f'{{{W_NS}}}t'):
                    if t.text:
                        if in_field_text:
                            field_text_parts.append(t.text)
                        else:
                            parts.append(t.text)

    return ''.join(parts).strip()


def _cell_para_to_md_text(para_element, bookmark_map, rel_map):
    """Convert a table cell paragraph to markdown text with checkboxes."""
    return _para_to_md_text_with_checkboxes(para_element, bookmark_map, rel_map)


# ─── CELL SHADING → RISK RATING ──────────────────────────────────────────────

# AUSTRAC uses colored cell fills instead of text for risk ratings.
# Map known fill colors to their rating labels.
SHADING_TO_RATING = {
    'BF4B3B': 'High',       # red
    'BF4B3F': 'High',       # red variant
    'F9B24D': 'Medium',     # amber
    'F9B24E': 'Medium',     # amber variant
    '68C3B5': 'Low',        # green
    '68C3B4': 'Low',        # green variant
}


def _cell_shading_rating(cell):
    """If a cell is empty but has a colored fill, return the rating label."""
    shd_elements = cell._element.findall(f'.//{{{W_NS}}}shd')
    for shd in shd_elements:
        fill = shd.get(f'{{{W_NS}}}fill', '')
        if fill and fill != 'auto':
            # Try exact match and case-insensitive
            rating = SHADING_TO_RATING.get(fill.upper())
            if not rating:
                rating = SHADING_TO_RATING.get(fill)
            if rating:
                return rating
    return None


# ─── TABLE CELL CONVERSION ──────────────────────────────────────────────────

def _build_style_map(doc):
    """Build a map from style IDs to style names."""
    style_map = {}
    for style in doc.styles:
        if style.style_id:
            style_map[style.style_id] = style.name
    return style_map


def _tc_shading_rating(tc_element):
    """If a tc element is empty but has a colored fill, return the rating."""
    shd_elements = tc_element.findall(f'.//{{{W_NS}}}shd')
    for shd in shd_elements:
        fill = shd.get(f'{{{W_NS}}}fill', '')
        if fill and fill != 'auto':
            rating = SHADING_TO_RATING.get(fill.upper())
            if not rating:
                rating = SHADING_TO_RATING.get(fill)
            if rating:
                return rating
    return None


def _tc_to_md_text(tc_element, bookmark_map, rel_map, style_map):
    """Convert a raw XML tc element to markdown text.

    Uses raw XML instead of python-docx Cell objects to avoid
    python-docx bugs with cell deduplication in irregular tables.
    """
    parts = []
    paras = tc_element.findall(f'{{{W_NS}}}p')
    for p_el in paras:
        text = _cell_para_to_md_text(p_el, bookmark_map, rel_map)
        if not text:
            continue
        # Get style name from pStyle
        pPr = p_el.find(f'{{{W_NS}}}pPr')
        style_id = ''
        if pPr is not None:
            pStyle = pPr.find(f'{{{W_NS}}}pStyle')
            if pStyle is not None:
                style_id = pStyle.get(f'{{{W_NS}}}val', '')
        style_name = style_map.get(style_id, style_id)
        style_lower = style_name.lower()
        # Check if paragraph has numId=0 (explicitly removed from list)
        is_deactivated = _has_deactivated_numbering(p_el)
        # 'Bullet list' items are sub-bullets (indented in Word)
        if style_lower == 'bullet list':
            parts.append('  - ' + text)
        elif _is_list_style(style_name) and not is_deactivated:
            parts.append('- ' + text)
        else:
            parts.append(text)

    # If cell is empty, check for color-coded risk rating
    if not parts:
        rating = _tc_shading_rating(tc_element)
        if rating:
            return rating

    return '<br>'.join(parts)


# ─── TABLE CONVERSION ───────────────────────────────────────────────────────

def table_to_md(table, bookmark_map, rel_map, style_map):
    """Convert a docx table to markdown table.

    Uses raw XML tc elements instead of python-docx Cell objects to
    correctly handle tables with irregular cell counts per row.
    """
    tbl = table._tbl

    # Get grid column count
    grid = tbl.find(f'{{{W_NS}}}tblGrid')
    grid_cols = len(grid.findall(f'{{{W_NS}}}gridCol')) if grid is not None else 0

    rows = []
    for tr in tbl.findall(f'{{{W_NS}}}tr'):
        # Collect tc elements: direct children OR inside sdt/sdtContent
        tcs = []
        for child in tr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tc':
                tcs.append(child)
            elif tag == 'sdt':
                content = child.find(f'{{{W_NS}}}sdtContent')
                if content is not None:
                    for sc in content:
                        sc_tag = sc.tag.split('}')[-1] if '}' in sc.tag else sc.tag
                        if sc_tag == 'tc':
                            tcs.append(sc)
        cells = []
        for tc in tcs:
            cell_text = _tc_to_md_text(tc, bookmark_map, rel_map, style_map)
            # Check gridSpan for cells spanning multiple columns
            tcPr = tc.find(f'{{{W_NS}}}tcPr')
            span = 1
            if tcPr is not None:
                gs = tcPr.find(f'{{{W_NS}}}gridSpan')
                if gs is not None:
                    span = int(gs.get(f'{{{W_NS}}}val', '1'))
            cells.append(cell_text)
            # Add empty cells for extra grid columns spanned
            for _ in range(span - 1):
                cells.append('')
        rows.append(cells)

    if not rows:
        return ''

    # Use header row or grid to determine column count
    num_cols = grid_cols if grid_cols > 0 else len(rows[0])

    # Standard markdown table
    lines = []
    lines.append('| ' + ' | '.join(rows[0][:num_cols]) + ' |')
    lines.append('| ' + ' | '.join([':---'] * num_cols) + ' |')
    for row in rows[1:]:
        while len(row) < num_cols:
            row.append('')
        lines.append('| ' + ' | '.join(row[:num_cols]) + ' |')

    return '\n'.join(lines)


# ─── PARAGRAPH STYLE DETECTION ──────────────────────────────────────────────

def _has_deactivated_numbering(para_element):
    """Check if a paragraph has numId=0, meaning explicitly removed from list."""
    pPr = para_element.find(f'{{{W_NS}}}pPr')
    if pPr is None:
        return False
    numPr = pPr.find(f'{{{W_NS}}}numPr')
    if numPr is None:
        return False
    numId_el = numPr.find(f'{{{W_NS}}}numId')
    if numId_el is not None and numId_el.get(f'{{{W_NS}}}val') == '0':
        return True
    return False


def _get_num_id(para_element):
    """Get the numId value for a paragraph, or None if not in a list."""
    pPr = para_element.find(f'{{{W_NS}}}pPr')
    if pPr is None:
        return None
    numPr = pPr.find(f'{{{W_NS}}}numPr')
    if numPr is None:
        return None
    numId_el = numPr.find(f'{{{W_NS}}}numId')
    if numId_el is not None:
        val = numId_el.get(f'{{{W_NS}}}val', '0')
        return int(val)
    return None


def _is_list_style(style_name):
    """Check if a paragraph style represents a list item."""
    s = style_name.lower()
    return 'list' in s or 'bullet' in s or 'number' in s


# ─── MAIN DOCUMENT CONVERSION ───────────────────────────────────────────────

def docx_to_md(docx_path):
    """Convert a .docx file to a markdown string."""
    doc = Document(docx_path)
    md_lines = []
    prev_was_list = False

    # Build lookup maps for cross-references, external links, and styles
    bookmark_map = _build_bookmark_map(doc)
    rel_map = _build_rel_map(doc)
    style_map = _build_style_map(doc)

    # Build ordered sequence of body elements (paragraphs and tables)
    table_index = 0
    para_index = 0
    body_elements = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            if para_index < len(doc.paragraphs):
                body_elements.append(('para', doc.paragraphs[para_index]))
                para_index += 1
        elif tag == 'tbl':
            if table_index < len(doc.tables):
                body_elements.append(('table', doc.tables[table_index]))
                table_index += 1

    numbered_counter = 0   # Tracks current number in a numbered list
    last_num_id = None     # Tracks numId to detect list continuity

    for elem_type, elem in body_elements:
        if elem_type == 'table':
            # End any open list before a table
            if prev_was_list:
                md_lines.append('')
                prev_was_list = False
            numbered_counter = 0
            last_num_id = None
            md_lines.append('')
            md_lines.append(table_to_md(elem, bookmark_map, rel_map, style_map))
            md_lines.append('')

        elif elem_type == 'para':
            p = elem
            style_name = p.style.name if p.style else ''
            style_lower = style_name.lower()

            # Get text with hyperlinks resolved to markdown links
            text = _para_to_md_text(p._element, bookmark_map, rel_map)

            # Empty paragraph — close list if open, otherwise just blank line
            if not text:
                if prev_was_list:
                    md_lines.append('')
                    prev_was_list = False
                elif md_lines and md_lines[-1] != '':
                    md_lines.append('')
                continue

            # Skip TOC entries (auto-generated, content is in the headings)
            if style_name.startswith('toc'):
                continue

            # Detect numbered list items (numId present and > 0)
            num_id = _get_num_id(p._element)
            is_numbered = (num_id is not None and num_id > 0
                           and _is_list_style(style_name))
            is_indented_bullet = 'increased indent' in style_lower
            is_list = _is_list_style(style_name)

            # Transition from list → non-list: insert blank line separator
            if prev_was_list and not is_list:
                md_lines.append('')

            # Headings
            if 'Heading 1' in style_name or style_name == 'Title':
                md_lines.append(f'# {text}')
                md_lines.append('')
                prev_was_list = False
                numbered_counter = 0
                last_num_id = None
            elif 'Heading 2' in style_name:
                md_lines.append(f'## {text}')
                md_lines.append('')
                prev_was_list = False
                numbered_counter = 0
                last_num_id = None
            elif 'Heading 3' in style_name:
                md_lines.append(f'### {text}')
                md_lines.append('')
                prev_was_list = False
                numbered_counter = 0
                last_num_id = None
            elif 'Heading 4' in style_name:
                md_lines.append(f'#### {text}')
                md_lines.append('')
                prev_was_list = False
                numbered_counter = 0
                last_num_id = None
            # Numbered list items
            elif is_numbered:
                numbered_counter += 1
                md_lines.append(f'{numbered_counter}. {text}')
                prev_was_list = True
            # Indented bullet items (sub-bullets under numbered items)
            elif is_indented_bullet:
                md_lines.append(f'   - {text}')
                prev_was_list = True
            # Regular bullet list items
            elif is_list:
                md_lines.append(f'- {text}')
                prev_was_list = True
            # Instructions (AUSTRAC guidance notes in italic blocks)
            elif style_name == 'Instructions' and text:
                md_lines.append(f'> *{text}*')
                md_lines.append('')
                prev_was_list = False
            # Normal text
            else:
                md_lines.append(text)
                md_lines.append('')
                prev_was_list = False

    # Clean up: collapse 3+ blank lines to 2
    result = '\n'.join(md_lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip() + '\n'


# ─── FILENAME CLEANING ──────────────────────────────────────────────────────

def clean_filename(docx_name, sector):
    """Convert docx filename to clean markdown filename.

    'Real estate - Customer forms - Onboarding form - Individual - January 2026.docx'
    → 'onboarding-form-individual.md'
    """
    name = docx_name.replace('.docx', '')

    # Remove sector prefix
    for prefix in ['Real estate - ', 'Accountants - ', 'Jewellers - ']:
        if name.startswith(prefix):
            name = name[len(prefix):]

    # Remove category prefix
    for prefix in ['Customer forms -  ', 'Customer forms - ', 'Customer forms -',
                    'Personnel forms - ', 'Maintain program forms - ']:
        if name.startswith(prefix):
            name = name[len(prefix):]

    # Remove date suffix
    name = re.sub(r'\s*-?\s*January 2026\s*', '', name)

    # Clean up to kebab-case
    name = name.strip(' -')
    name = re.sub(r'\s+', '-', name)
    name = re.sub(r'-+', '-', name)
    name = name.lower()

    return name + '.md'


# ─── MAIN ───────────────────────────────────────────────────────────────────

def main():
    # Single file mode: --file X.docx --out Y.md
    if '--file' in sys.argv:
        idx = sys.argv.index('--file')
        src = sys.argv[idx + 1]
        out_idx = sys.argv.index('--out') if '--out' in sys.argv else None
        out = sys.argv[out_idx + 1] if out_idx else src.replace('.docx', '.md')
        md = docx_to_md(src)
        with open(out, 'w', encoding='utf-8') as f:
            f.write(md)
        print(f'OK  {out} ({len(md)} bytes)')
        return

    # Batch mode: convert all sectors
    total = 0
    errors = 0

    for sector in SECTORS:
        for category in CATEGORIES:
            src_dir = os.path.join(BASE, sector, category)
            if not os.path.exists(src_dir):
                continue

            out_dir = os.path.join(BASE, 'markdown', sector, category)
            os.makedirs(out_dir, exist_ok=True)

            docx_files = [f for f in os.listdir(src_dir) if f.endswith('.docx')]
            for docx_file in sorted(docx_files):
                docx_path = os.path.join(src_dir, docx_file)
                md_name = clean_filename(docx_file, sector)
                md_path = os.path.join(out_dir, md_name)

                try:
                    md_content = docx_to_md(docx_path)
                    with open(md_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
                    total += 1
                    print(f'  OK  {sector}/{category}/{md_name}')
                except Exception as e:
                    errors += 1
                    print(f'  ERR {sector}/{category}/{docx_file}: {e}')

    print(f'\nConverted: {total}  Errors: {errors}')


if __name__ == '__main__':
    main()
